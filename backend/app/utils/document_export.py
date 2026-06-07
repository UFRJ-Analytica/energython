from __future__ import annotations

import base64
import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.+)$")
_ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")
_TABLE_RE = re.compile(r"^\s*\|(.+)\|\s*$")


def binary_to_base64(data: bytes) -> str:
    return base64.b64encode(data).decode("ascii")


def markdown_to_docx_base64(markdown: str) -> str:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.bold = True

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        if not raw.strip():
            i += 1
            continue

        if _TABLE_RE.match(raw):
            rows: list[list[str]] = []
            while i < len(lines) and _TABLE_RE.match(lines[i].rstrip()):
                cells = [clean_inline(cell.strip()) for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[:\-\s]+", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                table = document.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = cell_text
                        if r_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                document.add_paragraph()
            continue

        heading = _HEADING_RE.match(raw)
        if heading:
            level = min(len(heading.group(1)), 3)
            paragraph = document.add_heading(clean_inline(heading.group(2)), level=level)
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        bullet = _BULLET_RE.match(raw)
        ordered = _ORDERED_RE.match(raw)
        if bullet or ordered:
            style = "List Bullet" if bullet else "List Number"
            paragraph = document.add_paragraph(style=style)
            add_inline_runs(paragraph, (bullet or ordered).group(1))
            i += 1
            continue

        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, raw)
        i += 1

    buffer = BytesIO()
    document.save(buffer)
    return binary_to_base64(buffer.getvalue())


def markdown_to_pdf_base64(markdown: str) -> str:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title="Pleito de ressarcimento",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="DocTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=16, leading=20, alignment=1, spaceAfter=12))
    styles.add(ParagraphStyle(name="Section", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, spaceBefore=10, spaceAfter=6, textColor=colors.HexColor("#0f766e")))
    styles.add(ParagraphStyle(name="Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13, spaceAfter=5))
    styles.add(ParagraphStyle(name="BulletDoc", parent=styles["Body"], leftIndent=12, bulletIndent=4))

    story = []
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        if not raw.strip():
            i += 1
            continue

        if _TABLE_RE.match(raw):
            rows: list[list[str]] = []
            while i < len(lines) and _TABLE_RE.match(lines[i].rstrip()):
                cells = [pdf_escape(clean_inline(cell.strip())) for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[:\-\s]+", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                table = Table(rows, repeatRows=1)
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e6fffb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f766e")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ]))
                story += [table, Spacer(1, 8)]
            continue

        heading = _HEADING_RE.match(raw)
        if heading:
            level = len(heading.group(1))
            text = pdf_escape(clean_inline(heading.group(2)))
            story.append(Paragraph(text, styles["DocTitle" if level == 1 else "Section"]))
            i += 1
            continue

        bullet = _BULLET_RE.match(raw)
        ordered = _ORDERED_RE.match(raw)
        if bullet or ordered:
            story.append(Paragraph(pdf_escape(clean_inline((bullet or ordered).group(1))), styles["BulletDoc"], bulletText="•"))
            i += 1
            continue

        story.append(Paragraph(pdf_escape(clean_inline(raw)), styles["Body"]))
        i += 1

    doc.build(story)
    return binary_to_base64(buffer.getvalue())


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text.strip()


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(clean_inline(part))
        if part.startswith("**") and part.endswith("**"):
            run.bold = True


def pdf_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
